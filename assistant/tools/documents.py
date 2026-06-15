"""Document I/O tool (assistant.tools.documents).

Provides:

- :func:`read_pdf` -- extract text from a PDF file (all pages, joined with
  blank lines).
- :func:`read_docx` -- extract text from a DOCX file (all paragraphs,
  joined with newlines).
- :func:`create_docx` -- generate a new ``.docx`` file from plain text or a
  list of paragraph strings.

Errors
------
``FileNotFoundError`` is raised by the read functions when *file_path* does
not exist -- this is the standard Python idiom for a missing file and needs
no custom wrapper.

:class:`DocumentError` (a :class:`RuntimeError`) is raised for read/parse
failures on a corrupt or unreadable PDF/DOCX, and for write failures during
:func:`create_docx`. These are not retryable: re-parsing or re-writing the
same file would produce the same outcome.

Data-fidelity principle
------------------------
The text returned by :func:`read_pdf` / :func:`read_docx` is the raw,
unmodified output of the underlying parser -- no reformatting or
summarisation is performed here.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class DocumentError(RuntimeError):
    """Raised when a document cannot be read or written. Non-retryable."""


def read_pdf(file_path: str | Path) -> str:
    """Extract all text from a PDF file.

    Args:
        file_path: Path to the ``.pdf`` file.

    Returns:
        The extracted text, with each page's text separated by a blank
        line. Returns an empty string for a PDF with no extractable text.

    Raises:
        FileNotFoundError: If *file_path* does not exist.
        DocumentError: If the file cannot be parsed as a PDF.
    """
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        reader = PdfReader(str(path))
        pages_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise DocumentError(f"Unable to read PDF '{path}': {exc}") from exc

    return "\n\n".join(pages_text).strip()


def read_docx(file_path: str | Path) -> str:
    """Extract all text from a DOCX file.

    Args:
        file_path: Path to the ``.docx`` file.

    Returns:
        The extracted text, with each paragraph on its own line. Returns
        an empty string for a DOCX with no paragraphs.

    Raises:
        FileNotFoundError: If *file_path* does not exist.
        DocumentError: If the file cannot be parsed as a DOCX.
    """
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"DOCX file not found: {path}")

    try:
        document = Document(str(path))
        paragraphs = [p.text for p in document.paragraphs]
    except Exception as exc:
        raise DocumentError(f"Unable to read DOCX '{path}': {exc}") from exc

    return "\n".join(paragraphs).strip()


def create_docx(file_path: str | Path, content: str | list[str]) -> Path:
    """Create a new DOCX file containing *content*.

    Args:
        file_path: Destination path for the ``.docx`` file. Parent
            directories are created automatically if they do not exist.
        content: Either a single string (split into paragraphs on
            newlines) or a list of paragraph strings.

    Returns:
        The :class:`~pathlib.Path` the file was written to.

    Raises:
        ValueError: If *content* is empty, or whitespace-only.
        DocumentError: If the file cannot be written.
    """
    if isinstance(content, str):
        if not content.strip():
            raise ValueError("create_docx() requires non-empty content")
        paragraphs = content.splitlines()
    else:
        paragraphs = list(content)
        if not paragraphs or all(not p.strip() for p in paragraphs):
            raise ValueError("create_docx() requires non-empty content")

    path = Path(file_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    try:
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(str(path))
    except Exception as exc:
        raise DocumentError(f"Unable to create DOCX '{path}': {exc}") from exc

    logger.info("create_docx: wrote %d paragraph(s) to %s", len(paragraphs), path)
    return path
